from docx import Document
import sqlite3
import os

# ==========================================================
# DADOS FIXOS DO USUÁRIO
# ==========================================================

NOME = "Abimael Pereira De Moura"
EMAIL = "abimaelpereira297@gmail.com"
TELEFONE = "(51) 99625-1978"
CIDADE = "Cachoeirinha - RS"

RESUMO = """
Profissional com experiência em Controle de Qualidade, Inspeção de Produtos,
Conferência, Não Conformidades, Auditorias Internas, ISO 9001 e Melhoria Contínua.
Atuação em ambientes logísticos e industriais, com foco em processos, indicadores
e garantia da qualidade.
"""

EXPERIENCIAS = [
    {
        "empresa": "Magalog Serviços Logísticos LTDA",
        "cargo": "Assistente de Qualidade / Conferente",
        "periodo": "Mai/2021 - Jun/2025",
        "atividades": [
            "Inspeção e conferência de produtos",
            "Controle de não conformidades",
            "Acompanhamento de processos logísticos",
            "Elaboração de relatórios e indicadores",
            "Aplicação de melhoria contínua"
        ]
    }
]

FORMACAO = [
    "Tecnólogo em Gestão da Qualidade - UNINTER (Concluído)",
    "Graduação em Logística - UNINTER (2021 - 2023)"
]

COMPETENCIAS = [
    "ISO 9001",
    "Auditorias Internas",
    "Controle de Qualidade",
    "Inspeção de Produtos",
    "Não Conformidades",
    "5S",
    "Indicadores",
    "Melhoria Contínua",
    "Pacote Office",
    "Google Sheets"
]


# ==========================================================
# GERAR CURRÍCULO PERSONALIZADO
# ==========================================================

def gerar_curriculo():

    conn = sqlite3.connect("vagas.db")
    cursor = conn.cursor()

    # Busca a melhor vaga disponível pelo score final
    cursor.execute("""
        SELECT titulo, empresa, descricao, score_final
        FROM vagas
        ORDER BY score_final DESC
        LIMIT 1
    """)

    vaga = cursor.fetchone()
    conn.close()

    if not vaga:
        print("Nenhuma vaga encontrada no banco.")
        return

    titulo, empresa, descricao, score_final = vaga

    doc = Document()

    # Cabeçalho
    doc.add_heading(NOME, level=1)
    doc.add_paragraph(f"{CIDADE} | {TELEFONE} | {EMAIL}")

    # Resumo
    doc.add_heading("Resumo Profissional", level=2)
    doc.add_paragraph(RESUMO)

    # Experiência
    doc.add_heading("Experiência Profissional", level=2)

    for exp in EXPERIENCIAS:
        doc.add_heading(
            f"{exp['cargo']} - {exp['empresa']}",
            level=3
        )
        doc.add_paragraph(exp["periodo"])

        for atividade in exp["atividades"]:
            doc.add_paragraph(atividade, style="List Bullet")

    # Formação
    doc.add_heading("Formação Acadêmica", level=2)

    for curso in FORMACAO:
        doc.add_paragraph(curso, style="List Bullet")

    # Competências
    doc.add_heading("Competências", level=2)

    for comp in COMPETENCIAS:
        doc.add_paragraph(comp, style="List Bullet")

    # Adequação à vaga
    doc.add_heading("Adequação à Vaga", level=2)
    doc.add_paragraph(
        f"Este currículo foi adaptado automaticamente para a vaga '{titulo}' "
        f"da empresa {empresa}, que possui score final {score_final}, destacando "
        f"as experiências e competências mais alinhadas com a área de Qualidade."
    )

    # Cria a pasta
    os.makedirs("curriculos/docx", exist_ok=True)

    # Nome do arquivo
    nome_arquivo = (
        f"curriculos/docx/Curriculo_"
        f"{empresa.replace(' ', '')}_"
        f"{titulo.replace(' ', '')}.docx"
    )

    doc.save(nome_arquivo)

    print("=" * 60)
    print("CURRÍCULO GERADO COM SUCESSO")
    print("=" * 60)
    print(f"Vaga: {titulo}")
    print(f"Empresa: {empresa}")
    print(f"Score Final: {score_final}")
    print(f"Arquivo: {nome_arquivo}")
    print("=" * 60)


if __name__ == "__main__":
    gerar_curriculo()